"""
Comprehensive Report Builder for Customer Quality Report.

Generates rich HTML and DOCX reports with all ITIL process data.
"""

import base64
from datetime import datetime
from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from jinja2 import Template

from config import OUTPUT_DIR, COLORS
from analyzer import ComprehensiveResult, SLABreakdown
from i18n import get_all_texts


# =============================================================================
# HTML Template - Comprehensive Customer Quality Report
# =============================================================================

COMPREHENSIVE_HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="{{ language }}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ texts.report_title }}</title>
    <style>
        :root {
            /* 2B Business Professional Color Palette */
            --primary: #1e3a5f;
            --primary-light: #2c5282;
            --accent: #3182ce;
            --success: #276749;
            --success-light: #c6f6d5;
            --warning: #c05621;
            --warning-light: #feebc8;
            --danger: #9b2c2c;
            --danger-light: #fed7d7;
            --neutral: #718096;
            --neutral-light: #e2e8f0;

            /* Grays */
            --gray-50: #f7fafc;
            --gray-100: #edf2f7;
            --gray-200: #e2e8f0;
            --gray-300: #cbd5e0;
            --gray-400: #a0aec0;
            --gray-500: #718096;
            --gray-600: #4a5568;
            --gray-700: #2d3748;
            --gray-800: #1a202c;
            --gray-900: #171923;

            /* Status Colors */
            --status-normal: var(--success);
            --status-warning: var(--warning);
            --status-danger: var(--danger);
        }

        * { margin: 0; padding: 0; box-sizing: border-box; }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Microsoft YaHei', sans-serif;
            line-height: 1.6;
            color: var(--gray-700);
            background: var(--gray-50);
            font-size: 14px;
        }

        .container {
            max-width: 1240px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.08);
        }

        /* Header - Professional Solid Color */
        .report-header {
            background: var(--primary);
            color: white;
            padding: 1.75rem 2rem;
            border-bottom: 3px solid var(--accent);
        }
        .report-header h1 {
            font-size: 1.5rem;
            font-weight: 600;
            margin-bottom: 0.25rem;
            letter-spacing: -0.025em;
        }
        .report-meta {
            display: flex;
            flex-wrap: wrap;
            gap: 1.5rem;
            margin-top: 0.75rem;
            font-size: 0.8125rem;
            opacity: 0.85;
        }
        .report-meta-item {
            display: flex;
            align-items: center;
            gap: 0.375rem;
        }

        /* Navigation - Clean Minimal */
        .nav-bar {
            background: var(--gray-800);
            padding: 0 2rem;
            position: sticky;
            top: 0;
            z-index: 100;
            border-bottom: 1px solid var(--gray-700);
        }
        .nav-bar ul {
            display: flex;
            gap: 0;
            list-style: none;
        }
        .nav-bar li {
            border-right: 1px solid var(--gray-700);
        }
        .nav-bar li:first-child {
            border-left: 1px solid var(--gray-700);
        }
        .nav-bar a {
            display: block;
            color: var(--gray-300);
            text-decoration: none;
            font-size: 0.8125rem;
            font-weight: 500;
            padding: 0.625rem 1rem;
            transition: all 0.15s;
        }
        .nav-bar a:hover {
            color: white;
            background: var(--gray-700);
        }

        /* Main Content */
        .main-content {
            padding: 1.75rem 2rem;
        }

        /* Section */
        section {
            margin-bottom: 2rem;
        }
        .section-title {
            display: flex;
            align-items: center;
            gap: 0.625rem;
            font-size: 1rem;
            font-weight: 600;
            color: var(--gray-800);
            margin-bottom: 1rem;
            padding-bottom: 0.625rem;
            border-bottom: 1px solid var(--gray-200);
        }
        .section-title .icon {
            width: 22px;
            height: 22px;
            display: flex;
            align-items: center;
            justify-content: center;
            background: var(--primary);
            color: white;
            border-radius: 4px;
            font-size: 0.75rem;
            font-weight: 600;
        }
            font-size: 0.875rem;
        }

        /* Health Score - 2B Business Professional */
        .health-container {
            display: grid;
            grid-template-columns: 260px 1fr;
            gap: 1.5rem;
            margin-bottom: 1.5rem;
        }
        .health-score-card {
            background: var(--primary);
            color: white;
            padding: 1.5rem;
            border-radius: 4px;
            text-align: center;
            border: 1px solid var(--primary);
        }
        .health-score-card.at-risk {
            background: var(--danger);
            border-color: var(--danger);
        }
        .health-score-card.needs-improvement {
            background: var(--warning);
            border-color: var(--warning);
        }
        .health-score-card.good {
            background: #d69e2e;
            border-color: #d69e2e;
            color: white;
        }
        .health-score-card.excellent {
            background: var(--success);
            border-color: var(--success);
        }
        .health-score-value {
            font-size: 3rem;
            font-weight: 700;
            line-height: 1;
            font-family: 'Consolas', 'Monaco', monospace;
        }
        .health-score-label {
            font-size: 0.875rem;
            margin-top: 0.375rem;
            opacity: 0.8;
            font-weight: 500;
        }
        .health-score-grade {
            font-size: 1.125rem;
            font-weight: 600;
            margin-top: 0.625rem;
            letter-spacing: 0.025em;
        }

        .health-breakdown {
            display: grid;
            grid-template-columns: repeat(2, 1fr);
            gap: 0.75rem;
        }
        .health-factor {
            background: white;
            padding: 0.875rem 1rem;
            border-radius: 2px;
            border: 1px solid var(--gray-200);
            border-left: 3px solid var(--primary);
        }
        .health-factor.danger { border-left-color: var(--danger); }
        .health-factor.warning { border-left-color: var(--warning); }
        .health-factor.success { border-left-color: var(--success); }
        .health-factor-label {
            font-size: 0.6875rem;
            color: var(--gray-500);
            text-transform: uppercase;
            letter-spacing: 0.05em;
            font-weight: 500;
        }
        .health-factor-value {
            font-size: 1.25rem;
            font-weight: 600;
            color: var(--gray-800);
            font-family: 'Consolas', 'Monaco', monospace;
        }
        .health-factor-trend {
            font-size: 0.75rem;
            margin-top: 0.25rem;
        }
        .trend-up { color: var(--success); }
        .trend-down { color: var(--danger); }
        .trend-neutral { color: var(--gray-500); }

        /* KPI Grid - 2B Business Professional */
        .kpi-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
            gap: 0.75rem;
        }
        .kpi-card {
            background: white;
            border: 1px solid var(--gray-200);
            border-radius: 2px;
            padding: 1rem;
            border-left: 3px solid var(--gray-300);
        }
        .kpi-card.danger { border-left-color: var(--danger); background: #fef2f2; }
        .kpi-card.warning { border-left-color: var(--warning); background: #fffbeb; }
        .kpi-card.success { border-left-color: var(--success); background: #f0fdf4; }

        .kpi-header {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            margin-bottom: 0.375rem;
        }
        .kpi-label {
            font-size: 0.6875rem;
            color: var(--gray-500);
            text-transform: uppercase;
            letter-spacing: 0.05em;
            font-weight: 500;
        }
        .kpi-badge {
            font-size: 0.5625rem;
            padding: 0.125rem 0.375rem;
            border-radius: 2px;
            font-weight: 600;
            text-transform: uppercase;
        }
        .kpi-badge.incident { background: #1e40af; color: white; }
        .kpi-badge.change { background: #065f46; color: white; }
        .kpi-badge.request { background: #92400e; color: white; }
        .kpi-badge.problem { background: #9d174d; color: white; }

        .kpi-value {
            font-size: 1.5rem;
            font-weight: 600;
            color: var(--gray-800);
            line-height: 1;
            font-family: 'Consolas', 'Monaco', monospace;
        }
        .kpi-unit {
            font-size: 0.875rem;
            font-weight: 400;
            color: var(--gray-500);
        }
        .kpi-trend {
            display: flex;
            align-items: center;
            gap: 0.25rem;
            font-size: 0.75rem;
            margin-top: 0.375rem;
        }

        /* Tables - 2B Business Professional */
        .data-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 0.8125rem;
            border: 1px solid var(--gray-200);
        }
        .data-table th {
            background: var(--gray-800);
            color: white;
            padding: 0.625rem 0.75rem;
            text-align: left;
            font-weight: 500;
            font-size: 0.75rem;
            text-transform: uppercase;
            letter-spacing: 0.025em;
            border-bottom: none;
        }
        .data-table td {
            padding: 0.5rem 0.75rem;
            border-bottom: 1px solid var(--gray-200);
            vertical-align: middle;
        }
        .data-table tr:nth-child(even) td {
            background: var(--gray-50);
        }
        .data-table tr:hover td {
            background: #f0f7ff;
        }
        .data-table .priority-badge {
            display: inline-block;
            padding: 0.125rem 0.5rem;
            border-radius: 2px;
            font-size: 0.6875rem;
            font-weight: 600;
            min-width: 28px;
            text-align: center;
        }
        .priority-p1 { background: var(--danger); color: white; }
        .priority-p2 { background: var(--warning); color: white; }
        .priority-p3 { background: var(--accent); color: white; }
        .priority-p4 { background: var(--gray-500); color: white; }

        /* Risk Cards - 2B Business Professional */
        .risk-grid {
            display: grid;
            gap: 0.625rem;
        }
        .risk-card {
            display: flex;
            gap: 0.75rem;
            padding: 0.75rem 1rem;
            border-radius: 2px;
            border-left: 3px solid;
            background: white;
            border: 1px solid var(--gray-200);
        }
        .risk-critical {
            background: #fef2f2;
            border-left: 3px solid var(--danger);
        }
        .risk-warning {
            background: #fffbeb;
            border-left: 3px solid var(--warning);
        }
        .risk-attention {
            background: #eff6ff;
            border-left: 3px solid var(--accent);
        }
        .risk-icon {
            width: 28px;
            height: 28px;
            display: flex;
            align-items: center;
            justify-content: center;
            border-radius: 2px;
            font-weight: 600;
            font-size: 0.75rem;
            flex-shrink: 0;
        }
        .risk-critical .risk-icon { background: var(--danger); color: white; }
        .risk-warning .risk-icon { background: var(--warning); color: white; }
        .risk-attention .risk-icon { background: var(--accent); color: white; }
        .risk-content { flex: 1; min-width: 0; }
        .risk-title {
            font-weight: 600;
            color: var(--gray-800);
            margin-bottom: 0.125rem;
            font-size: 0.8125rem;
        }
        .risk-impact {
            font-size: 0.75rem;
            color: var(--gray-600);
        }
        .risk-process {
            font-size: 0.5625rem;
            padding: 0.125rem 0.375rem;
            border-radius: 2px;
            background: var(--gray-200);
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: var(--gray-600);
            font-weight: 500;
        }

        /* Charts Section - 2B Business Professional */
        .chart-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 1rem;
        }
        .chart-card {
            background: white;
            border: 1px solid var(--gray-200);
            border-radius: 2px;
            padding: 1rem;
        }
        .chart-title {
            font-size: 0.75rem;
            font-weight: 600;
            color: var(--gray-700);
            margin-bottom: 0.75rem;
            text-transform: uppercase;
            letter-spacing: 0.025em;
            padding-bottom: 0.5rem;
            border-bottom: 1px solid var(--gray-200);
        }
        .chart-container {
            text-align: center;
        }
        .chart-container img {
            max-width: 100%;
            border-radius: 2px;
        }

        /* Action Items - 2B Business Professional */
        .action-list {
            display: grid;
            gap: 0.5rem;
        }
        .action-item {
            display: flex;
            align-items: flex-start;
            gap: 0.75rem;
            padding: 0.75rem 1rem;
            background: white;
            border: 1px solid var(--gray-200);
            border-radius: 2px;
        }
        .action-priority {
            padding: 0.125rem 0.5rem;
            border-radius: 2px;
            font-size: 0.625rem;
            font-weight: 600;
            text-transform: uppercase;
            white-space: nowrap;
        }
        .action-urgent { background: var(--danger); color: white; }
        .action-high { background: var(--warning); color: white; }
        .action-medium { background: var(--primary); color: white; }
        .action-content { flex: 1; }
        .action-title {
            font-weight: 600;
            color: var(--gray-800);
            font-size: 0.8125rem;
        }
        .action-impact {
            font-size: 0.75rem;
            color: var(--gray-600);
            margin-top: 0.125rem;
        }
        .action-process {
            font-size: 0.5625rem;
            color: var(--gray-500);
            text-transform: uppercase;
        }

        /* Process Summary Cards - 2B Business Professional */
        .process-grid {
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 0.75rem;
        }
        .process-card {
            background: white;
            border: 1px solid var(--gray-200);
            border-radius: 2px;
            padding: 1rem;
            border-top: 3px solid;
        }
        .process-incident { border-top-color: #1e40af; }
        .process-change { border-top-color: #065f46; }
        .process-request { border-top-color: #92400e; }
        .process-problem { border-top-color: #9d174d; }

        .process-title {
            font-size: 0.6875rem;
            font-weight: 600;
            color: var(--gray-500);
            margin-bottom: 0.375rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }
        .process-count {
            font-size: 1.75rem;
            font-weight: 600;
            color: var(--gray-800);
            font-family: 'Consolas', 'Monaco', monospace;
        }
        .process-label {
            font-size: 0.6875rem;
            color: var(--gray-500);
        }
        .process-stats {
            margin-top: 0.75rem;
            padding-top: 0.75rem;
            border-top: 1px solid var(--gray-200);
        }
        .process-stat {
            display: flex;
            justify-content: space-between;
            font-size: 0.75rem;
            margin-bottom: 0.375rem;
        }
        .process-stat-label { color: var(--gray-600); }
        .process-stat-value { font-weight: 600; color: var(--gray-800); font-family: 'Consolas', 'Monaco', monospace; }

        /* Footer - 2B Business Professional */
        .report-footer {
            background: var(--gray-800);
            color: var(--gray-300);
            padding: 1rem 2rem;
            display: flex;
            justify-content: space-between;
            align-items: center;
            font-size: 0.75rem;
        }
        .footer-meta {
            display: flex;
            gap: 1.5rem;
        }

        /* Responsive */
        @media (max-width: 1024px) {
            .health-container { grid-template-columns: 1fr; }
            .process-grid { grid-template-columns: repeat(2, 1fr); }
        }
        @media (max-width: 640px) {
            .kpi-grid { grid-template-columns: repeat(2, 1fr); }
            .process-grid { grid-template-columns: 1fr; }
        }

        /* Print styles */
        @media print {
            .nav-bar { display: none; }
            body { background: white; }
            .container { box-shadow: none; }
        }
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <header class="report-header">
            <h1>{{ texts.report_title }}</h1>
            <div class="report-meta">
                <div class="report-meta-item">
                    <span>{{ texts.period }}:</span>
                    <strong>{{ result.start_date }} {{ texts.to }} {{ result.end_date }}</strong>
                </div>
                <div class="report-meta-item">
                    <span>{{ texts.data_span }}:</span>
                    <strong>{{ result.data_span_days }} days</strong>
                </div>
                <div class="report-meta-item">
                    <span>Report ID:</span>
                    <strong>CQR-{{ result.start_date | replace('-', '') }}-{{ result.end_date | replace('-', '') }}</strong>
                </div>
            </div>
        </header>

        <!-- Navigation -->
        <nav class="nav-bar">
            <ul>
                <li><a href="#health">Health Score</a></li>
                <li><a href="#overview">Process Overview</a></li>
                <li><a href="#kpis">KPI Dashboard</a></li>
                <li><a href="#sla">SLA Details</a></li>
                <li><a href="#risks">Risk Radar</a></li>
                <li><a href="#actions">Actions</a></li>
                <li><a href="#details">Details</a></li>
            </ul>
        </nav>

        <main class="main-content">
            <!-- Health Score Section -->
            <section id="health">
                <h2 class="section-title">
                    <span class="icon">H</span>
                    {{ texts.section_health_score }}
                </h2>
                <div class="health-container">
                    <div class="health-score-card {{ health_class }}">
                        <div class="health-score-value">{{ result.health_score | round | int }}</div>
                        <div class="health-score-label">/ 100</div>
                        <div class="health-score-grade">{{ result.health_emoji }} {{ result.health_grade }}</div>
                    </div>
                    <div class="health-breakdown">
                        {% for key, kpi in result.kpis.items() %}
                        {% if key in ['sla_rate', 'change_success_rate', 'request_csat', 'problem_closure_rate'] %}
                        <div class="health-factor {{ kpi.status }}">
                            <div class="health-factor-label">{{ kpi.name }}</div>
                            <div class="health-factor-value">
                                {% if 'rate' in key or 'sla' in key %}{{ "%.1f" | format(kpi.current_value * 100) }}%
                                {% elif 'csat' in key %}{{ "%.2f" | format(kpi.current_value) }}/5
                                {% else %}{{ kpi.current_value | int }}{% endif %}
                            </div>
                            {% if kpi.change is not none %}
                            <div class="health-factor-trend {% if kpi.trend == '↑' %}trend-up{% elif kpi.trend == '↓' %}trend-down{% else %}trend-neutral{% endif %}">
                                {{ kpi.trend }} {{ "%.1f" | format(kpi.change | abs * 100) }}% vs previous
                            </div>
                            {% endif %}
                        </div>
                        {% endif %}
                        {% endfor %}
                    </div>
                </div>
            </section>

            <!-- AI Insight (Executive) -->
            {% if insights.executive_summary %}
            <section id="insight">
                <h2 class="section-title">
                    <span class="icon">💡</span>
                    {{ texts.get('section_ai_insight', 'AI Insight') }}
                </h2>
                <div style="background: var(--primary-light); color: white; padding: 1.5rem; border-radius: 4px; border-left: 4px solid var(--accent); white-space: pre-wrap;">{{ insights.executive_summary }}</div>
            </section>
            {% endif %}

            <!-- Process Overview -->
            <section id="overview">
                <h2 class="section-title">
                    <span class="icon">P</span>
                    {{ texts.section_process_overview }}
                </h2>
                <div class="process-grid">
                    <!-- Incidents -->
                    <div class="process-card process-incident">
                        <div class="process-title">Incident Management</div>
                        <div class="process-count">{{ result.total_incidents }}</div>
                        <div class="process-label">Total Incidents</div>
                        <div class="process-stats">
                            {% if result.incident_summary and result.incident_summary.kpis %}
                            <div class="process-stat">
                                <span class="process-stat-label">SLA Rate</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.incident_summary.kpis.sla_rate.current_value * 100) }}%</span>
                            </div>
                            <div class="process-stat">
                                <span class="process-stat-label">Avg MTTR</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.incident_summary.kpis.avg_mttr.current_value) }}h</span>
                            </div>
                            {% endif %}
                        </div>
                    </div>

                    <!-- Changes -->
                    <div class="process-card process-change">
                        <div class="process-title">Change Management</div>
                        <div class="process-count">{{ result.total_changes }}</div>
                        <div class="process-label">Total Changes</div>
                        <div class="process-stats">
                            {% if result.change_summary and result.change_summary.kpis and result.change_summary.kpis.change_success_rate %}
                            <div class="process-stat">
                                <span class="process-stat-label">Success Rate</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.change_summary.kpis.change_success_rate.current_value * 100) }}%</span>
                            </div>
                            <div class="process-stat">
                                <span class="process-stat-label">Incidents Caused</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.change_summary.kpis.change_incident_rate.current_value * 100) }}%</span>
                            </div>
                            {% else %}
                            <div class="process-stat">
                                <span class="process-stat-label">No data</span>
                            </div>
                            {% endif %}
                        </div>
                    </div>

                    <!-- Requests -->
                    <div class="process-card process-request">
                        <div class="process-title">Service Requests</div>
                        <div class="process-count">{{ result.total_requests }}</div>
                        <div class="process-label">Total Requests</div>
                        <div class="process-stats">
                            {% if result.request_summary and result.request_summary.kpis and result.request_summary.kpis.request_sla_rate %}
                            <div class="process-stat">
                                <span class="process-stat-label">SLA Rate</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.request_summary.kpis.request_sla_rate.current_value * 100) }}%</span>
                            </div>
                            <div class="process-stat">
                                <span class="process-stat-label">CSAT</span>
                                <span class="process-stat-value">{{ "%.2f" | format(result.request_summary.kpis.request_csat.current_value) }}/5</span>
                            </div>
                            {% else %}
                            <div class="process-stat">
                                <span class="process-stat-label">No data</span>
                            </div>
                            {% endif %}
                        </div>
                    </div>

                    <!-- Problems -->
                    <div class="process-card process-problem">
                        <div class="process-title">Problem Management</div>
                        <div class="process-count">{{ result.total_problems }}</div>
                        <div class="process-label">Total Problems</div>
                        <div class="process-stats">
                            {% if result.problem_summary and result.problem_summary.kpis and result.problem_summary.kpis.problem_closure_rate %}
                            <div class="process-stat">
                                <span class="process-stat-label">Closure Rate</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.problem_summary.kpis.problem_closure_rate.current_value * 100) }}%</span>
                            </div>
                            <div class="process-stat">
                                <span class="process-stat-label">RCA Rate</span>
                                <span class="process-stat-value">{{ "%.1f" | format(result.problem_summary.kpis.rca_rate.current_value * 100) }}%</span>
                            </div>
                            {% else %}
                            <div class="process-stat">
                                <span class="process-stat-label">No data</span>
                            </div>
                            {% endif %}
                        </div>
                    </div>
                </div>
            </section>

            <!-- KPI Dashboard -->
            <section id="kpis">
                <h2 class="section-title">
                    <span class="icon">K</span>
                    {{ texts.section_kpi_dashboard }}
                </h2>
                <div class="kpi-grid">
                    {% for key, kpi in result.kpis.items() %}
                    <div class="kpi-card {{ kpi.status }}">
                        <div class="kpi-header">
                            <span class="kpi-label">{{ kpi.name }}</span>
                            <span class="kpi-badge {{ kpi.category }}">{{ kpi.category }}</span>
                        </div>
                        <div class="kpi-value">
                            {% if 'rate' in key or 'sla' in key %}{{ "%.1f" | format(kpi.current_value * 100) }}<span class="kpi-unit">%</span>
                            {% elif 'mttr' in key or 'time' in key %}{{ "%.1f" | format(kpi.current_value) }}<span class="kpi-unit">h</span>
                            {% elif 'csat' in key %}{{ "%.2f" | format(kpi.current_value) }}<span class="kpi-unit">/5</span>
                            {% else %}{{ kpi.current_value | int }}{% endif %}
                        </div>
                        {% if kpi.change is not none %}
                        <div class="kpi-trend {% if kpi.trend == '↑' %}trend-up{% elif kpi.trend == '↓' %}trend-down{% else %}trend-neutral{% endif %}">
                            {{ kpi.trend }} {{ "%.1f" | format(kpi.change | abs * 100) }}%
                        </div>
                        {% endif %}
                    </div>
                    {% endfor %}
                </div>
            </section>

            <!-- SLA Breakdown -->
            {% if result.sla_breakdown %}
            <section id="sla">
                <h2 class="section-title">
                    <span class="icon">S</span>
                    {{ texts.section_sla_breakdown }}
                </h2>
                <table class="data-table">
                    <thead>
                        <tr>
                            <th>Priority</th>
                            <th>Target (hours)</th>
                            <th>Total Tickets</th>
                            <th>Compliant</th>
                            <th>SLA Rate</th>
                            <th>Status</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for sla in result.sla_breakdown %}
                        <tr>
                            <td><span class="priority-badge priority-{{ sla.priority | lower }}">{{ sla.priority }}</span></td>
                            <td>{{ sla.target }}h</td>
                            <td>{{ sla.total }}</td>
                            <td>{{ sla.compliant }}</td>
                            <td><strong>{{ "%.1f" | format(sla.rate * 100) }}%</strong></td>
                            <td>
                                {% if sla.rate >= 0.90 %}
                                <span style="color: var(--success);">✓ On Track</span>
                                {% elif sla.rate >= 0.80 %}
                                <span style="color: var(--warning);">⚠ At Risk</span>
                                {% else %}
                                <span style="color: var(--danger);">✗ Below Target</span>
                                {% endif %}
                            </td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </section>
            {% endif %}

            <!-- Risk Radar -->
            {% if result.top_risks %}
            <section id="risks">
                <h2 class="section-title">
                    <span class="icon">R</span>
                    {{ texts.section_risk_radar }}
                </h2>
                {% if insights.risk_insight %}
                <div style="background: var(--warning-light); padding: 1rem; border-radius: 4px; margin-bottom: 1rem; white-space: pre-wrap; font-size: 0.9em; border: 1px solid var(--warning);"><strong>AI Analysis:</strong> {{ insights.risk_insight }}</div>
                {% endif %}
                <div class="risk-grid">
                    {% for risk in result.top_risks %}
                    <div class="risk-card risk-{{ risk.priority | lower }}">
                        <div class="risk-icon">{{ risk.id[1:] }}</div>
                        <div class="risk-content">
                            <div class="risk-title">{{ risk.message }}</div>
                            <div class="risk-impact">{{ risk.impact }}</div>
                        </div>
                        <span class="risk-process">{{ risk.process }}</span>
                    </div>
                    {% endfor %}
                </div>
            </section>
            {% endif %}

            <!-- Trends -->
            {% if charts %}
            <section id="trends">
                <h2 class="section-title">
                    <span class="icon">T</span>
                    {{ texts.section_trends }}
                </h2>
                <div class="chart-grid">
                    {% if charts.incident_volume %}
                    <div class="chart-card">
                        <div class="chart-title">Incident Volume Trend</div>
                        <div class="chart-container">
                            <img src="data:image/png;base64,{{ charts.incident_volume }}" alt="Incident Volume">
                        </div>
                    </div>
                    {% endif %}
                    {% if charts.sla_trend %}
                    <div class="chart-card">
                        <div class="chart-title">SLA Compliance Trend</div>
                        <div class="chart-container">
                            <img src="data:image/png;base64,{{ charts.sla_trend }}" alt="SLA Trend">
                        </div>
                    </div>
                    {% endif %}
                    {% if charts.mttr_trend %}
                    <div class="chart-card">
                        <div class="chart-title">MTTR Trend</div>
                        <div class="chart-container">
                            <img src="data:image/png;base64,{{ charts.mttr_trend }}" alt="MTTR Trend">
                        </div>
                    </div>
                    {% endif %}
                </div>
            </section>
            {% endif %}

            <!-- Recommended Actions -->
            {% if result.actions %}
            <section id="actions">
                <h2 class="section-title">
                    <span class="icon">A</span>
                    {{ texts.section_actions }}
                </h2>
                {% if insights.action_insight %}
                <div style="background: var(--gray-100); padding: 1rem; border-radius: 4px; margin-bottom: 1rem; white-space: pre-wrap; font-size: 0.9em; border: 1px solid var(--gray-300);"><strong>AI Analysis:</strong> {{ insights.action_insight }}</div>
                {% endif %}
                <div class="action-list">
                    {% for action in result.actions %}
                    <div class="action-item">
                        <span class="action-priority action-{{ action.priority | lower }}">{{ action.priority }}</span>
                        <div class="action-content">
                            <div class="action-title">{{ action.action }}</div>
                            <div class="action-impact">Expected: {{ action.expected_impact }}</div>
                            <div class="action-process">{{ action.process }}</div>
                        </div>
                    </div>
                    {% endfor %}
                </div>
            </section>
            {% endif %}

            <!-- Major Incidents -->
            {% if result.major_incidents %}
            <section id="details">
                <h2 class="section-title">
                    <span class="icon">M</span>
                    {{ texts.section_major_incidents }}
                </h2>
                <table class="data-table">
                    <thead>
                        <tr>
                            <th>ID</th>
                            <th>Description</th>
                            <th>Priority</th>
                            <th>Status</th>
                            <th>Resolution Time</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for inc in result.major_incidents %}
                        <tr>
                            <td><code>{{ inc.order_number }}</code></td>
                            <td>{{ inc.name[:60] }}{% if inc.name | length > 60 %}...{% endif %}</td>
                            <td><span class="priority-badge priority-{{ inc.priority | lower }}">{{ inc.priority }}</span></td>
                            <td>{{ inc.status }}</td>
                            <td>{% if inc.resolution_time %}{{ "%.1f" | format(inc.resolution_time / 60) }}h{% else %}—{% endif %}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </section>
            {% endif %}

            <!-- Failed Changes -->
            {% if result.failed_changes %}
            <section>
                <h2 class="section-title">
                    <span class="icon">F</span>
                    {{ texts.section_failed_changes }}
                </h2>
                <table class="data-table">
                    <thead>
                        <tr>
                            <th>Change ID</th>
                            <th>Title</th>
                            <th>Type</th>
                            <th>Status</th>
                            <th>Incident Caused</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for chg in result.failed_changes %}
                        <tr>
                            <td><code>{{ chg.change_number }}</code></td>
                            <td>{{ chg.title }}</td>
                            <td>{{ chg.change_type }}</td>
                            <td>{{ chg.status }}</td>
                            <td>{% if chg.incident_caused %}<span style="color: var(--danger);">Yes</span>{% else %}No{% endif %}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </section>
            {% endif %}

            <!-- Open Problems -->
            {% if result.open_problems %}
            <section>
                <h2 class="section-title">
                    <span class="icon">O</span>
                    {{ texts.section_open_problems }}
                </h2>
                <table class="data-table">
                    <thead>
                        <tr>
                            <th>Problem ID</th>
                            <th>Title</th>
                            <th>Status</th>
                            <th>Known Error</th>
                            <th>Related Incidents</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for prb in result.open_problems %}
                        <tr>
                            <td><code>{{ prb.problem_number }}</code></td>
                            <td>{{ prb.title }}</td>
                            <td>{{ prb.status }}</td>
                            <td>{% if prb.known_error %}Yes{% else %}No{% endif %}</td>
                            <td>{{ prb.related_incidents }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </section>
            {% endif %}

        </main>

        <!-- Footer -->
        <footer class="report-footer">
            <div class="footer-meta">
                <span>{{ texts.generated_on }}: {{ generated_time }}</span>
                <span>Classification: Internal Use Only</span>
            </div>
            <div>
                <span>Customer Quality Report v2.0</span>
            </div>
        </footer>
    </div>
</body>
</html>
"""


class ComprehensiveReportBuilder:
    """Report builder for comprehensive customer quality reports."""

    def __init__(
        self,
        result: ComprehensiveResult,
        charts: Dict[str, str] = None,
        insights: Dict[str, str] = None,
        language: str = "en"
    ):
        self.result = result
        self.charts = charts or {}
        self.insights = insights or {}
        self.language = language
        self.texts = get_all_texts(language)

        # Add new text keys if not present
        self._extend_texts()

    def _extend_texts(self):
        """Extend texts with additional keys."""
        defaults = {
            "section_process_overview": "Process Overview" if self.language == "en" else "流程概览",
            "section_sla_breakdown": "SLA Breakdown by Priority" if self.language == "en" else "按优先级的 SLA 明细",
            "section_failed_changes": "Failed Changes" if self.language == "en" else "失败的变更",
            "section_open_problems": "Open Problems" if self.language == "en" else "未关闭的问题",
            "section_major_incidents": "Major Incidents (P1/P2)" if self.language == "en" else "重大事件 (P1/P2)",
            "data_span": "Data Span" if self.language == "en" else "数据跨度",
            "section_ai_insight": "AI Insight" if self.language == "en" else "AI 智能洞察",
        }
        for key, value in defaults.items():
            if key not in self.texts:
                self.texts[key] = value

    def build_html(self) -> str:
        """Build comprehensive HTML report."""
        template = Template(COMPREHENSIVE_HTML_TEMPLATE)

        # Determine health class
        if self.result.health_score >= 90:
            health_class = "excellent"
        elif self.result.health_score >= 80:
            health_class = "good"
        elif self.result.health_score >= 70:
            health_class = "needs-improvement"
        else:
            health_class = "at-risk"

        return template.render(
            language=self.language,
            texts=self.texts,
            result=self.result,
            charts=self.charts,
            insights=self.insights,
            health_class=health_class,
            generated_time=datetime.now().strftime("%Y-%m-%d %H:%M")
        )

    def build_docx(self) -> Document:
        """Build comprehensive DOCX report with professional styling."""
        from io import BytesIO
        import base64
        from office_theme import (
            COLORS, hex_to_rgb, set_cell_shading, set_cell_border,
            get_status_color, get_status_bg_color, get_priority_color,
            get_process_color, FONTS, FONT_SIZES
        )

        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # =================================================================
        # HEADER - Colored title bar
        # =================================================================
        header_table = doc.add_table(rows=1, cols=1)
        header_table.autofit = False
        header_cell = header_table.rows[0].cells[0]
        set_cell_shading(header_cell, COLORS["primary"])

        # Title
        p = header_cell.paragraphs[0]
        title_run = p.add_run(self.texts["report_title"])
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = hex_to_rgb(COLORS["text_light"])

        # Subtitle (period)
        p = header_cell.add_paragraph()
        sub_run = p.add_run(
            f"{self.texts['period']}: {self.result.start_date} {self.texts['to']} {self.result.end_date}"
        )
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = hex_to_rgb(COLORS["gray_300"])

        # Set cell padding
        header_cell.paragraphs[0].paragraph_format.space_before = Pt(12)
        header_cell.paragraphs[-1].paragraph_format.space_after = Pt(12)
        for para in header_cell.paragraphs:
            para.paragraph_format.left_indent = Inches(0.15)

        doc.add_paragraph()

        # =================================================================
        # HEALTH SCORE SECTION
        # =================================================================
        doc.add_heading(self.texts["section_health_score"], 1)

        # Determine health color
        score = self.result.health_score
        if score >= 90:
            health_color = COLORS["success"]
            health_bg = COLORS["success_light"]
        elif score >= 80:
            health_color = COLORS["warning"]
            health_bg = COLORS["warning_light"]
        else:
            health_color = COLORS["danger"]
            health_bg = COLORS["danger_light"]

        # Health score box using table
        health_table = doc.add_table(rows=1, cols=2)
        health_table.autofit = False

        # Score cell
        score_cell = health_table.rows[0].cells[0]
        score_cell.width = Inches(2)
        set_cell_shading(score_cell, health_color)

        p = score_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = p.add_run(f"{score:.0f}")
        score_run.bold = True
        score_run.font.size = Pt(48)
        score_run.font.color.rgb = hex_to_rgb(COLORS["text_light"])

        p = score_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        grade_run = p.add_run(f"{self.result.health_emoji} {self.result.health_grade}")
        grade_run.font.size = Pt(12)
        grade_run.font.color.rgb = hex_to_rgb(COLORS["text_light"])

        # Key metrics cell
        metrics_cell = health_table.rows[0].cells[1]
        set_cell_shading(metrics_cell, COLORS["gray_100"])
        metrics_cell.paragraphs[0].paragraph_format.left_indent = Inches(0.15)

        key_kpis = ['sla_rate', 'change_success_rate', 'request_csat', 'problem_closure_rate']
        for kpi_key in key_kpis:
            if kpi_key in self.result.kpis:
                kpi = self.result.kpis[kpi_key]
                p = metrics_cell.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

                # KPI name
                name_run = p.add_run(f"{kpi.name}: ")
                name_run.font.size = Pt(10)
                name_run.font.color.rgb = hex_to_rgb(COLORS["text_secondary"])

                # KPI value
                if "rate" in kpi_key or "sla" in kpi_key:
                    value_text = f"{kpi.current_value:.1%}"
                elif "csat" in kpi_key:
                    value_text = f"{kpi.current_value:.2f}/5"
                else:
                    value_text = f"{kpi.current_value:.1f}"

                value_run = p.add_run(value_text)
                value_run.bold = True
                value_run.font.size = Pt(11)
                value_run.font.color.rgb = hex_to_rgb(get_status_color(kpi.status))

                # Trend
                if kpi.change is not None:
                    trend_run = p.add_run(f" ({kpi.trend}{abs(kpi.change):.1%})")
                    trend_run.font.size = Pt(9)
                    trend_run.font.color.rgb = hex_to_rgb(COLORS["gray_500"])

        doc.add_paragraph()

        # =================================================================
        # PROCESS OVERVIEW
        # =================================================================
        doc.add_heading(self.texts["section_process_overview"], 1)

        process_table = doc.add_table(rows=2, cols=4)
        process_table.autofit = False

        process_data = [
            ("Incidents" if self.language == "en" else "事件", self.result.total_incidents, COLORS["incident"]),
            ("Changes" if self.language == "en" else "变更", self.result.total_changes, COLORS["change"]),
            ("Requests" if self.language == "en" else "请求", self.result.total_requests, COLORS["request"]),
            ("Problems" if self.language == "en" else "问题", self.result.total_problems, COLORS["problem"]),
        ]

        for i, (name, count, color) in enumerate(process_data):
            # Header cell
            header_cell = process_table.rows[0].cells[i]
            set_cell_shading(header_cell, color)
            p = header_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = hex_to_rgb(COLORS["text_light"])

            # Value cell
            value_cell = process_table.rows[1].cells[i]
            set_cell_shading(value_cell, COLORS["gray_50"])
            p = value_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(count))
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = hex_to_rgb(color)

        doc.add_paragraph()

        # =================================================================
        # KPI DASHBOARD
        # =================================================================
        doc.add_heading(self.texts["section_kpi_dashboard"], 1)
        self._add_styled_kpi_table(doc)
        doc.add_paragraph()

        # AI Executive Insight
        if self.insights.get("executive_summary"):
            doc.add_heading(self.texts.get('section_ai_insight', 'AI Insight'), 2)
            
            # Container table for styling
            insight_table = doc.add_table(rows=1, cols=1)
            insight_table.autofit = False
            cell = insight_table.rows[0].cells[0]
            set_cell_shading(cell, COLORS["primary_light"]) # Dark blue background
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add icon/title
            title_run = p.add_run(f"💡 {self.texts.get('section_ai_insight', 'AI Insight')}\n")
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = hex_to_rgb(COLORS["text_light"])
            
            # Add insight text
            text_run = p.add_run(self.insights["executive_summary"])
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = hex_to_rgb(COLORS["text_light"])
            
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.right_indent = Inches(0.1)

            doc.add_paragraph()

        # =================================================================
        # SLA BREAKDOWN
        # =================================================================
        if self.result.sla_breakdown:
            doc.add_heading(self.texts["section_sla_breakdown"], 1)
            sla_table = doc.add_table(rows=len(self.result.sla_breakdown) + 1, cols=5)

            # Header row
            headers = [
                self.texts.get("sla_priority", "Priority"),
                self.texts.get("sla_target", "Target"),
                self.texts.get("sla_total", "Total"),
                self.texts.get("sla_compliant", "Compliant"),
                self.texts.get("sla_rate", "SLA Rate"),
            ]
            for i, h in enumerate(headers):
                cell = sla_table.rows[0].cells[i]
                set_cell_shading(cell, COLORS["gray_800"])
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = hex_to_rgb(COLORS["text_light"])

            # Data rows
            for j, sla in enumerate(self.result.sla_breakdown):
                row = sla_table.rows[j + 1]

                # Determine status color
                if sla.rate >= 0.90:
                    status_color = COLORS["success"]
                    row_bg = COLORS["success_light"]
                elif sla.rate >= 0.80:
                    status_color = COLORS["warning"]
                    row_bg = COLORS["warning_light"]
                else:
                    status_color = COLORS["danger"]
                    row_bg = COLORS["danger_light"]

                data = [sla.priority, f"{sla.target}h", str(sla.total), str(sla.compliant), f"{sla.rate:.1%}"]
                for k, val in enumerate(data):
                    cell = row.cells[k]
                    if k == 4:  # Rate column gets colored
                        set_cell_shading(cell, row_bg)
                    p = cell.paragraphs[0]
                    run = p.add_run(val)
                    run.font.size = Pt(10)
                    if k == 4:
                        run.bold = True
                        run.font.color.rgb = hex_to_rgb(status_color)

            doc.add_paragraph()

        # =================================================================
        # EMBEDDED CHARTS
        # =================================================================
        if self.charts:
            doc.add_heading(self.texts["section_trends"], 1)

            chart_keys = ['incident_volume', 'sla_trend', 'mttr_trend', 'health_gauge']
            for chart_key in chart_keys:
                if chart_key in self.charts and self.charts[chart_key]:
                    try:
                        img_data = base64.b64decode(self.charts[chart_key])
                        img_stream = BytesIO(img_data)
                        doc.add_picture(img_stream, width=Inches(5.5))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                    except Exception:
                        pass

        # =================================================================
        # RISK RADAR
        # =================================================================
        if self.result.top_risks:
            doc.add_heading(self.texts["section_risk_radar"], 1)

            # AI Risk Insight
            if self.insights.get("risk_insight"):
                p = doc.add_paragraph()
                run = p.add_run(f"AI Analysis: {self.insights['risk_insight']}")
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgb(COLORS["warning"])
                p.paragraph_format.space_after = Pt(12)

            for risk in self.result.top_risks:
                # Risk card using table
                risk_table = doc.add_table(rows=1, cols=1)
                risk_cell = risk_table.rows[0].cells[0]

                # Set background based on priority
                priority_color = get_priority_color(risk.priority)
                set_cell_shading(risk_cell, get_status_bg_color(risk.priority))
                set_cell_border(risk_cell, priority_color, "12")

                # Priority badge
                p = risk_cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(8)
                badge_run = p.add_run(f"[{risk.priority.upper()}] ")
                badge_run.bold = True
                badge_run.font.size = Pt(9)
                badge_run.font.color.rgb = hex_to_rgb(priority_color)

                # Message
                msg_run = p.add_run(risk.message)
                msg_run.font.size = Pt(11)
                msg_run.font.color.rgb = hex_to_rgb(COLORS["text_primary"])

                # Impact
                p = risk_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                impact_run = p.add_run(f"Impact: {risk.impact}")
                impact_run.font.size = Pt(9)
                impact_run.font.color.rgb = hex_to_rgb(COLORS["text_secondary"])
                impact_run.italic = True

                doc.add_paragraph()

        # =================================================================
        # RECOMMENDED ACTIONS
        # =================================================================
        if self.result.actions:
            doc.add_heading(self.texts["section_actions"], 1)

            # AI Action Insight
            if self.insights.get("action_insight"):
                p = doc.add_paragraph()
                run = p.add_run(f"AI Analysis: {self.insights['action_insight']}")
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgb(COLORS["text_secondary"])
                p.paragraph_format.space_after = Pt(12)

            for action in self.result.actions:
                action_table = doc.add_table(rows=1, cols=1)
                action_cell = action_table.rows[0].cells[0]
                set_cell_shading(action_cell, COLORS["gray_100"])

                p = action_cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(8)

                # Priority badge
                priority_color = get_priority_color(action.priority)
                badge_run = p.add_run(f"[{action.priority}] ")
                badge_run.bold = True
                badge_run.font.size = Pt(9)
                badge_run.font.color.rgb = hex_to_rgb(priority_color)

                # Action
                action_run = p.add_run(action.action)
                action_run.font.size = Pt(11)
                action_run.font.color.rgb = hex_to_rgb(COLORS["text_primary"])

                # Expected impact
                p = action_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                impact_run = p.add_run(f"Expected: {action.expected_impact}")
                impact_run.font.size = Pt(9)
                impact_run.font.color.rgb = hex_to_rgb(COLORS["text_secondary"])

                doc.add_paragraph()

        # =================================================================
        # MAJOR INCIDENTS
        # =================================================================
        if self.result.major_incidents:
            doc.add_heading(self.texts["section_major_incidents"], 1)
            inc_table = doc.add_table(rows=min(len(self.result.major_incidents), 10) + 1, cols=4)

            # Header
            headers = ["ID", "Description", "Priority", "Resolution"]
            for i, h in enumerate(headers):
                cell = inc_table.rows[0].cells[i]
                set_cell_shading(cell, COLORS["incident"])
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = hex_to_rgb(COLORS["text_light"])

            # Data
            for j, inc in enumerate(self.result.major_incidents[:10]):
                row = inc_table.rows[j + 1]
                res_time = f"{inc.resolution_time/60:.1f}h" if inc.resolution_time else "-"
                data = [inc.order_number, inc.name[:50], inc.priority, res_time]
                for k, val in enumerate(data):
                    run = row.cells[k].paragraphs[0].add_run(val)
                    run.font.size = Pt(9)

            doc.add_paragraph()

        # =================================================================
        # FOOTER
        # =================================================================
        doc.add_paragraph()
        footer_table = doc.add_table(rows=1, cols=1)
        footer_cell = footer_table.rows[0].cells[0]
        set_cell_shading(footer_cell, COLORS["gray_200"])

        p = footer_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = f"{self.texts['generated_on']}: {datetime.now().strftime('%Y-%m-%d %H:%M')} | {self.texts.get('footer_confidential', 'Confidential - Internal Use Only')}"
        run = p.add_run(footer_text)
        run.font.size = Pt(8)
        run.font.color.rgb = hex_to_rgb(COLORS["text_secondary"])

        return doc

    def _add_styled_kpi_table(self, doc: Document) -> None:
        """Add styled KPI table with color coding."""
        from office_theme import COLORS, hex_to_rgb, set_cell_shading, get_status_color

        kpis = list(self.result.kpis.values())
        cols = min(len(kpis), 4)
        rows = (len(kpis) + cols - 1) // cols

        table = doc.add_table(rows=rows * 2, cols=cols)

        for i, kpi in enumerate(kpis):
            row_idx = (i // cols) * 2
            col_idx = i % cols

            # Label cell
            label_cell = table.rows[row_idx].cells[col_idx]
            set_cell_shading(label_cell, COLORS["gray_100"])
            p = label_cell.paragraphs[0]
            run = p.add_run(kpi.name)
            run.font.size = Pt(9)
            run.font.color.rgb = hex_to_rgb(COLORS["text_secondary"])

            # Value cell
            value_cell = table.rows[row_idx + 1].cells[col_idx]

            # Format value based on type
            if "rate" in kpi.name.lower() or "sla" in kpi.name.lower():
                value_text = f"{kpi.current_value:.1%}"
            elif "mttr" in kpi.name.lower() or "time" in kpi.name.lower():
                value_text = f"{kpi.current_value:.1f}h"
            elif "csat" in kpi.name.lower():
                value_text = f"{kpi.current_value:.2f}/5"
            else:
                value_text = f"{int(kpi.current_value)}"

            # Apply status color
            status_color = get_status_color(kpi.status)
            p = value_cell.paragraphs[0]
            run = p.add_run(value_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = hex_to_rgb(status_color)

            # Trend
            if kpi.change is not None:
                trend_run = p.add_run(f" {kpi.trend}{abs(kpi.change):.1%}")
                trend_run.font.size = Pt(9)
                trend_run.font.color.rgb = hex_to_rgb(COLORS["gray_500"])

    def save(self, base_filename: str = None) -> Dict[str, str]:
        """Save reports to files."""
        if base_filename is None:
            lang = "CN" if self.language == "zh" else "EN"
            base_filename = f"Operation_Quality_Report_{self.result.start_date}_to_{self.result.end_date}_{lang}"

        paths = {}

        # HTML
        html_path = OUTPUT_DIR / f"{base_filename}.html"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(self.build_html())
        paths["html"] = str(html_path)

        # DOCX
        docx_path = OUTPUT_DIR / f"{base_filename}.docx"
        self.build_docx().save(docx_path)
        paths["docx"] = str(docx_path)

        return paths
